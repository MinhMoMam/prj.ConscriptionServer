import pandas as pd
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from spire.doc import Document
from spire.doc import FileFormat
from spire.doc import BookmarksNavigator
import numbers
import decimal
import numpy
import yaml

class wordFileGenerator:
    def __init__(self, generatorConfigPath, lableConfigPath):
        # Load configration
        with open(lableConfigPath, "r", encoding="utf-8") as f:
            self.lableConfig = yaml.safe_load(f)
        with open(generatorConfigPath, "r", encoding="utf-8") as f:
            self.generatorConfig = yaml.safe_load(f)
        # Create lable Dictionary
        self.columnConverter = {}
        for column in self.lableConfig["columnList"]:
            key = self.lableConfig[column]["ColumnLabel"].strip()
            self.columnConverter[key] = column
        # Scan through all lables
        self.scanForAllLables()

    def create_doc_from_template(self):
        doc = Document()
        doc.LoadFromFile(self.generatorConfig["templateLocation"][0],FileFormat.Auto)
        return doc

    def scanForAllLables(self):
        doc = self.create_doc_from_template()
        self.labelList = []
        allText = doc.GetText()
        left = 0
        right = 0
        for i in range(len(allText)):
            if allText[i] == "<":
                left = i
            elif allText[i] == ">":
                right = i
            if right > left:
                self.labelList.append(allText[left:right+1])
                left = right

    def generateWordFile(self,input,outputPath):
        tempDoc = self.create_doc_from_template()
        for lables in self.labelList:
            parts = lables[1:-1].split(".")
            try:
                columnTile = self.columnConverter[parts[0].strip().replace("–", "-").replace("\\n", "\n")]
                data = input[columnTile]
            except:
                data = "<" + parts[0] + " - Not found>"
            if len(parts) != 1:
                if parts[1] == "Upper()":
                    data = data.upper()
                elif parts[1] == "Title()":
                    data = data.title()
                elif parts[1].startswith("Add"):
                    arg = parts[1].replace("Add","").replace("(","").replace(")","")
                    try:
                        data = str(int(data) + int(arg))
                    except:
                        print("[ERROR]: Invalid data type.")
                elif parts[1].startswith("Day()"):
                    date = data.split("/")
                    data = date[0]
                elif parts[1].startswith("Month()"):
                    date = data.split("/")
                    data = date[1]
                elif parts[1].startswith("Year()"):
                    date = data.split("/")
                    data = date[2]
                else:
                    print("[WARNING]: Unsuppored feature")
            tempDoc.Replace(lables,str(data),False, True)
        tempDoc.SaveToFile(outputPath, FileFormat.Docx2016)
        self.removeWaltermark(outputPath)


    def removeWaltermark(self,wordFileName):
        watermark = "Evaluation Warning: The document was created with Spire.Doc for Python."
        wordFile = docx.Document(wordFileName)
        for p in wordFile.paragraphs:
            for run in p.runs:
                if watermark in run.text:
                    run.text = run.text.replace(watermark,"")
                    par = p._element
                    par.getparent().remove(par)
                    par._p = par._element = None
        wordFile.save(wordFileName)